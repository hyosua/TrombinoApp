import sys
import os
from PIL import Image, ImageOps
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 1. Traitement d'Images (Pillow) ---

SUPPORTED_FORMATS = ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff')

def resize_image(input_path, output_dir, max_size_kb=200):
    """
    Redimensionne une image pour qu'elle pèse moins de max_size_kb.
    Gère plusieurs formats et préserve la transparence (PNG).
    """
    try:
        # Créer le dossier de sortie s'il n'existe pas
        os.makedirs(output_dir, exist_ok=True)
        
        # Définir le chemin de sortie
        filename = os.path.basename(input_path)
        base_name, ext = os.path.splitext(filename)
        # Nous convertissons en JPEG pour un contrôle de qualité, sauf si c'est un PNG
        output_format = 'PNG' if ext.lower() == '.png' else 'JPEG'
        output_ext = '.png' if output_format == 'PNG' else '.jpg'
        output_path = os.path.join(output_dir, f"{base_name}_processed{output_ext}")
        # Définir les dimensions cibles pour le "crop" (format carré)
        TARGET_DIMENSIONS = (300, 300)

        with Image.open(input_path) as img:
            # Corriger l'orientation EXIF si présente
            img = ImageOps.exif_transpose(img)

            # On rogne l'image par le centre pour qu'elle s'adapte
            # parfaitement aux dimensions cibles (ex: 300x300)
            img = ImageOps.fit(
                img, 
                TARGET_DIMENSIONS, 
                Image.LANCZOS, 
                centering=(0.5, 0.5) # Centrer le crop
            )
            
            # Si l'image a un canal Alpha (transparence), la garder en PNG
            if img.mode in ('RGBA', 'LA') or 'transparency' in img.info:
                output_format = 'PNG'
                output_ext = '.png'
                output_path = os.path.join(output_dir, f"{base_name}_processed.png")
            else:
                # Convertir en RGB si nécessaire (pour JPEG)
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                output_format = 'JPEG'
                output_ext = '.jpg'
                output_path = os.path.join(output_dir, f"{base_name}_processed.jpg")

            # Réduire la résolution si l'image est très grande
            img.thumbnail((1024, 1024), Image.LANCZOS)

            # Logique pour atteindre la taille cible
            quality = 90
            while quality > 10:
                # Sauvegarder dans un buffer temporaire pour vérifier la taille
                import io
                buffer = io.BytesIO()
                if output_format == 'JPEG':
                    img.save(buffer, format=output_format, quality=quality, optimize=True)
                else:
                    img.save(buffer, format=output_format, optimize=True) # PNG
                
                size_kb = buffer.tell() / 1024
                
                if size_kb <= max_size_kb:
                    # C'est bon, on sauvegarde sur le disque
                    with open(output_path, 'wb') as f:
                        f.write(buffer.getvalue())
                    return output_path
                
                # Réduire la qualité pour le prochain essai
                quality -= 10
            
            # Si on n'y arrive pas (très rare), sauvegarder la version la plus basse
            with open(output_path, 'wb') as f:
                f.write(buffer.getvalue())
            return output_path

    except Exception as e:
        print(f"Erreur redimensionnement {input_path}: {e}")
        return None

# --- 2. Lecteur Excel (openpyxl) ---

def read_excel(filepath):
    """
    Ouvre un fichier Excel et lit la première colonne (A) comme liste de noms.
    """
    try:
        workbook = load_workbook(filename=filepath, read_only=True)
        sheet = workbook.active
        
        student_list = []
        for row in sheet.iter_rows(min_row=1, max_col=1, values_only=True):
            if row[0]: # Ignorer les cellules vides
                student_list.append(str(row[0]).strip())
                
        return sorted(student_list) # Assurer le tri alphabétique
    except Exception as e:
        print(f"Erreur lecture Excel {filepath}: {e}")
        return None

# --- 3. Exportateur Word (python-docx) ---

def create_word_doc(associations, layout_str, save_path):
    """
    Crée un document Word .docx avec les photos et les noms.
    'associations' est un dict: {photo_path: student_name}
    'layout_str' est "3x4", "4x5", etc.
    """
    try:
        cols, rows_per_page = map(int, layout_str.split('x'))
        
        doc = Document()
        # Mettre des marges plus petites
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)

        table = doc.add_table(rows=1, cols=cols)
        table.autofit = False
        
        # Calcul de la largeur de colonne
        col_width = int((doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin) / cols)
        for col in table.columns:
            col.width = col_width

        row_cells = table.rows[0].cells
        item_index = 0
        
        # Trier les associations par nom d'étudiant pour l'export
        sorted_items = sorted(associations.items(), key=lambda item: item[1])
        total_items = len(sorted_items)

        for photo_path, student_name in sorted_items:
            cell = row_cells[item_index % cols]
            
            # Ajouter l'image
            try:
                run = cell.paragraphs[0].add_run()
                run.add_picture(photo_path, width=col_width * 0.9) # 90% de la largeur de cellule
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as img_e:
                cell.paragraphs[0].add_run(f"[Image {photo_path} illisible]")
                print(f"Erreur ajout image {photo_path} au DOCX: {img_e}")
            
            # Ajouter le nom
            p = cell.add_paragraph(student_name)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(10)
            
            item_index += 1
            
            # Ajouter une nouvelle ligne si nécessaire
            if item_index % cols == 0 and item_index < total_items:
                row_cells = table.add_row().cells
                for col in table.columns:
                    col.width = col_width # Réappliquer la largeur

        doc.save(save_path)
        return True
    except Exception as e:
        print(f"Erreur création DOCX: {e}")
        return False